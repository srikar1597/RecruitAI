"""
File text extraction utility.
Supports: PDF, DOCX, DOC, TXT
"""

import os


def extract_text(filepath: str) -> str:
    """Extract plain text from a resume file."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.txt':
        return _read_txt(filepath)
    elif ext == '.pdf':
        return _read_pdf(filepath)
    elif ext in ('.docx', '.doc'):
        return _read_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _read_txt(filepath: str) -> str:
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def _read_pdf(filepath: str) -> str:
    """Try pdfplumber first, fallback to PyPDF2."""
    text = ''

    # Method 1: pdfplumber (better quality)
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            pages = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            text = '\n'.join(pages)
        if text.strip():
            return text
    except Exception:
        pass

    # Method 2: PyPDF2 fallback
    try:
        import PyPDF2
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            text = '\n'.join(pages)
        if text.strip():
            return text
    except Exception:
        pass

    raise ValueError("Could not extract text from PDF. It may be a scanned document without OCR.")


def _read_docx(filepath: str) -> str:
    """Extract text from DOCX file."""
    try:
        import docx
        doc = docx.Document(filepath)
        parts = []

        # Main paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return '\n'.join(parts)
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {str(e)}")
